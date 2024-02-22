import pandas as pd
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pdf2docx import Converter
from docx2pdf import convert
import os
import asyncio

# Lendo a planilha do Excel
df = pd.read_excel('planilha.xlsx')

#Config server email via SMTP
server = smtplib.SMTP('smtp.gmail.com', 587)
server.starttls()
server.login('seuEmail@gmail.com', 'suaSenha') #email e senha do server ttls

def gerarPDF(name):
    doc = Document('basecertificate.docx')
    print(doc)
    # Substitua a palavra-chave
    for p in doc.paragraphs:
        if 'NOME' in p.text:
            p.text = p.text.replace('NOME', '')
            new_run = p.add_run(name)
            new_run.bold = True
            new_run.font.name = "Carlito"
            new_run.font.size = Pt(36)
            
            # Excluir o texto original (palavra-chave)
            for run in p.runs:
                if 'NOME' in run.text:
                    p.runs.remove(run)
                    break

    # Salve o documento com as alterações
    doc.save('certificado.docx')
    # Converta o documento Word em PDF
    convert('certificado.docx')


# Enviando o e-mail para cada linha da planilha
def RotinaDeEmails():
    for index, row in df.iterrows():
        if os.path.exists('certificado.pdf'):
            os.remove('certificado.pdf')
        if(row['Endereço de e-mail'] != '-------------------------------------------------------' and not pd.isna(row['Endereço de e-mail']) ):
            msg = MIMEMultipart()
            msg['From'] = 'seu_email@gmail.com' #email que envia o certificado
            msg['To'] = row['Endereço de e-mail']
            msg['Subject'] = '| Simpósio de liderança PMESP' #Assunto do email


            # Corpo do email
            corpo_do_email = 'Adicionar corpo do email'
            msg.attach(MIMEText(corpo_do_email, 'plain')) #Se quiser um corpo mais interessante mudar plain para html e adicionar teste HTML


            gerarPDF(row['Nome completo'])
            print('O Pdf foi gerado corretamente')


            # Anexando o certificado
            attachment = open('certificado.pdf', 'rb')
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment; filename= "certificado.pdf"')
            msg.attach(part)
            attachment.close()
            # Enviando o e-mail
            email = row['Endereço de e-mail']
            print(f"Iniciando envio do email para '{email}'")
            text = msg.as_string()
            server.sendmail('seuEmail@gmail.com', row['Endereço de e-mail'], text) 
            print("Email enviado com sucesso, passando para o próximo valor")
           
            
  
RotinaDeEmails()# Fechando a conexão com o servidor de e-mail
print("Processo finalizado")
server.quit()



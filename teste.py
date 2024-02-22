from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pdf2docx import Converter
from docx2pdf import convert

# Carregue o documento
doc = Document('certificado.docx')
print(doc)
# Substitua a palavra-chave
for p in doc.paragraphs:
    if 'GEOVANA DONELLA' in p.text:
        p.text = p.text.replace('GEOVANA DONELLA', '')
        new_run = p.add_run('GUSTAVO MOLINO')
        new_run.bold = True
        new_run.font.name = "Carlito"
        new_run.font.size = Pt(36)
        
        # Excluir o texto original (palavra-chave)
        for run in p.runs:
            if 'GEOVANA DONELLA' in run.text:
                p.runs.remove(run)
                break

# Salve o documento com as alterações
doc.save('documento_alterado.docx')
doc = Document('documento_alterado.docx')
print(doc)
# Converta o documento Word em PDF
convert('documento_alterado.docx')
